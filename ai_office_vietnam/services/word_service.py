from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

from ai_office_vietnam.config import AppSettings
from ai_office_vietnam.models import PageContent

TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*$")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


class WordService:
    def create_document(self, pages: list[PageContent], output_path: Path, settings: AppSettings,
                        images_by_page: dict[int, list[tuple[bytes, str]]] | None = None) -> None:
        doc = Document()
        section = doc.sections[0]
        section.page_width, section.page_height = Mm(210), Mm(297)
        section.left_margin, section.right_margin = Cm(settings.margin_left_cm), Cm(settings.margin_right_cm)
        section.top_margin, section.bottom_margin = Cm(settings.margin_top_cm), Cm(settings.margin_bottom_cm)
        normal = doc.styles["Normal"]
        normal.font.name, normal.font.size = settings.font_name, Pt(settings.font_size)
        normal.paragraph_format.line_spacing = settings.line_spacing
        normal.paragraph_format.space_after = Pt(0)
        images_by_page = images_by_page or {}
        ordered = sorted(pages, key=lambda p: p.page_number)
        for index, page in enumerate(ordered):
            self._append_markdown(doc, page.markdown)
            for image_bytes, _ext in images_by_page.get(page.page_number, []):
                try:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.add_run().add_picture(io.BytesIO(image_bytes), width=Cm(14.5))
                except Exception:
                    pass
            if index < len(ordered) - 1:
                doc.add_page_break()
        self._apply_font(doc, settings)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    def _append_markdown(self, doc: Document, markdown: str) -> None:
        lines = markdown.replace("\r", "").split("\n")
        i = 0
        while i < len(lines):
            raw = lines[i].strip()
            if not raw:
                i += 1
                continue
            if "|" in raw and i + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[i + 1].strip()):
                rows = []
                while i < len(lines) and "|" in lines[i] and lines[i].strip():
                    if not TABLE_SEPARATOR_RE.match(lines[i].strip()):
                        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                    i += 1
                if rows:
                    table = doc.add_table(rows=len(rows), cols=max(map(len, rows)))
                    table.style = "Table Grid"
                    for r, row in enumerate(rows):
                        for c, value in enumerate(row):
                            table.cell(r, c).text = value
                            if r == 0:
                                for run in table.cell(r, c).paragraphs[0].runs:
                                    run.bold = True
                continue
            alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style = None
            text = raw
            for start, end, align, sty in (
                ("[TITLE]", "[/TITLE]", WD_ALIGN_PARAGRAPH.CENTER, "Title"),
                ("[CENTER]", "[/CENTER]", WD_ALIGN_PARAGRAPH.CENTER, None),
                ("[RIGHT]", "[/RIGHT]", WD_ALIGN_PARAGRAPH.RIGHT, None),
            ):
                if raw.startswith(start) and raw.endswith(end):
                    text, alignment, style = raw[len(start):-len(end)].strip(), align, sty
                    break
            paragraph = doc.add_paragraph(style=style)
            paragraph.alignment = alignment
            pos = 0
            for match in INLINE_RE.finditer(text):
                paragraph.add_run(text[pos:match.start()])
                token = match.group(0)
                run = paragraph.add_run(token[2:-2] if token.startswith("**") else token[1:-1])
                run.bold = token.startswith("**")
                run.italic = not token.startswith("**")
                pos = match.end()
            paragraph.add_run(text[pos:])
            i += 1

    @staticmethod
    def _apply_font(doc: Document, settings: AppSettings) -> None:
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        for paragraph in paragraphs:
            paragraph.paragraph_format.line_spacing = settings.line_spacing
            paragraph.paragraph_format.space_after = Pt(0)
            size = settings.font_size + (2 if paragraph.style and paragraph.style.name == "Title" else 0)
            for run in paragraph.runs:
                run.font.name, run.font.size = settings.font_name, Pt(size)
                rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rfonts.set(qn(key), settings.font_name)
